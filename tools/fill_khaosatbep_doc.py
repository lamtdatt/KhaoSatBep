from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


DOC_PATH = Path(r"D:\ThucTapHM\Bai1\KhaoSatBep\BaoCao_KhaoSatBep.docx")
OUT_PATH = Path(r"D:\ThucTapHM\Bai1\KhaoSatBep\BaoCao_KhaoSatBep_DaDienThongTin_v2.docx")


def set_cell(cell, text):
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(11)


def set_para_text(paragraph, text, bold=False):
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run.bold = bold


def insert_paragraph_after(paragraph, text="", style=None, bold=False):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    new_para._element = new_p
    if style:
        new_para.style = style
    set_para_text(new_para, text, bold=bold)
    return new_para


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def bold_row(row):
    for cell in row.cells:
        shade_cell(cell, "D9EAF7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def fill_table(table, rows):
    # Reuse existing rows where possible, append if the new content is longer.
    while len(table.rows) < len(rows):
        table.add_row()

    # Clear extra old rows by leaving them blank if any remain.
    for r_idx, row in enumerate(table.rows):
        values = rows[r_idx] if r_idx < len(rows) else [""] * len(row.cells)
        for c_idx, cell in enumerate(row.cells):
            set_cell(cell, values[c_idx] if c_idx < len(values) else "")

    bold_row(table.rows[0])


def main():
    doc = Document(DOC_PATH)

    # Header information table.
    info = doc.tables[0]
    set_cell(info.cell(0, 1), "Ban Giám đốc Bệnh viện Hoàn Mỹ Đồng Nai")
    set_cell(info.cell(1, 1), 'Phê duyệt kế hoạch xây dựng Dự án cải tiến "Hệ thống Khảo Sát Bếp".')
    set_cell(info.cell(2, 1), "Bùi Thanh Hiền - Thực tập sinh Quản lý chất lượng")
    set_cell(info.cell(3, 1), "Ban Giám đốc xem xét và phê duyệt")
    set_cell(info.cell(4, 1), "16/05/2026")
    set_cell(info.cell(5, 1), "    /KH-QLCL")

    # Section 1 goals.
    set_para_text(doc.paragraphs[6], "1.1 Mục tiêu", bold=True)
    goals = [
        "- Số hóa quy trình lập biên bản kiểm tra bếp ăn, giảm thao tác ghi chép thủ công và giúp dữ liệu được lưu trữ tập trung.",
        "- Hỗ trợ nhân viên lập, gửi và theo dõi trạng thái các biên bản kiểm tra; hỗ trợ Admin tiếp nhận, xem xét và duyệt biên bản.",
        "- Cung cấp dashboard thống kê số lượng biên bản, trạng thái xử lý và tỷ lệ tiêu chí đạt/không đạt để phục vụ công tác quản lý chất lượng.",
    ]
    for idx, text in enumerate(goals, start=7):
        set_para_text(doc.paragraphs[idx], text)

    anchor = doc.paragraphs[9]
    additions = [
        ("1.2 Phạm vi thực hiện", True),
        ("- Xây dựng hệ thống web gồm giao diện nhân viên, giao diện Admin, API backend và cơ sở dữ liệu lưu biên bản.", False),
        ("- Áp dụng cho 4 loại biên bản: Cơ sở hạ tầng, Hồ sơ sổ sách, Vệ sinh an toàn thực phẩm và Suất ăn người bệnh.", False),
        ("1.3 Công nghệ sử dụng", True),
        ("- Frontend: Vue.js 3, Vite, Vue Router, Chart.js, HTML/CSS/JavaScript.", False),
        ("- Backend: ASP.NET Core Web API .NET 8, Entity Framework Core, JWT Authentication, BCrypt, Swagger.", False),
        ("- Cơ sở dữ liệu và triển khai: PostgreSQL/Supabase, Render cho API, Vercel cho frontend.", False),
        ("1.4 Chức năng chính", True),
        ("- Đăng nhập, phân quyền Admin/Nhân viên bằng JWT.", False),
        ("- Tạo, gửi, xem chi tiết, cập nhật trạng thái và duyệt biên bản.", False),
        ("- Quản lý chữ ký điện tử, thông báo trạng thái, thống kê theo loại biên bản và theo thời gian.", False),
    ]
    current = anchor
    for text, bold in additions:
        current = insert_paragraph_after(current, text, bold=bold)

    # Project execution plan.
    fill_table(doc.tables[1], [
        ["STT", "Kế hoạch hành động", "Đầu vào", "Mô tả chi tiết hành động", "Đầu ra", "Người làm hành động", "Người hỗ trợ phối hợp", "Thời hạn hoàn thành", "Giám sát thực hiện"],
        ["1", "GĐ 1: Khảo sát nghiệp vụ và thiết kế dữ liệu", "Yêu cầu kiểm tra bếp, mẫu biên bản hiện có", "Phân tích quy trình lập biên bản, xác định vai trò Admin/Nhân viên và thiết kế các bảng dữ liệu chính.", "Mô hình dữ liệu và danh sách chức năng", "Bùi Thanh Hiền", "Tổ IT / QLCL", "16/05", "TP QLCL và Tổ IT"],
        ["2", "GĐ 2: Xây dựng Backend API", "ASP.NET Core, PostgreSQL/Supabase", "Lập trình API đăng nhập, phân quyền JWT, CRUD biên bản, duyệt biên bản và thống kê.", "API KhaoSatBep hoạt động, có Swagger", "Bùi Thanh Hiền", "Tổ IT / QLCL", "20/05", "TP QLCL và Tổ IT"],
        ["3", "GĐ 3: Phát triển giao diện nhân viên", "Vue.js, Vite, các mẫu biên bản", "Xây dựng dashboard nhân viên, form nhập 4 loại biên bản, chữ ký điện tử và chức năng gửi biên bản.", "Giao diện nhân viên hoàn chỉnh", "Bùi Thanh Hiền", "Tổ IT / QLCL", "23/05", "TP QLCL và Tổ IT"],
        ["4", "GĐ 4: Phát triển giao diện Admin", "Dữ liệu biên bản từ API", "Xây dựng dashboard Admin, danh sách biên bản, xem chi tiết, duyệt biên bản, quản lý mẫu và thống kê.", "Giao diện Admin hoàn chỉnh", "Bùi Thanh Hiền", "Tổ IT / QLCL", "25/05", "TP QLCL và Tổ IT"],
        ["5", "GĐ 5: Kiểm thử và hoàn thiện triển khai", "Toàn bộ source code", "Kiểm thử luồng đăng nhập, tạo biên bản, gửi biên bản, duyệt biên bản, thống kê và cấu hình deploy.", "Sản phẩm hoàn chỉnh, sẵn sàng demo", "Bùi Thanh Hiền", "Tổ IT / QLCL", "26/05", "TP QLCL và Tổ IT"],
    ])

    # Progress report table.
    fill_table(doc.tables[2], [
        ["Giai đoạn", "Buổi", "Nội dung đạt được", "Hướng đi tiếp theo", "Người báo cáo", "Người rà soát"],
        ["GĐ 1: Khảo sát & CSDL", "B1 | 16/05", "Xác định bài toán quản lý biên bản kiểm tra bếp ăn; thống nhất 4 nhóm biên bản cần số hóa.", "Thiết kế database và cấu trúc API.", "Bùi Thanh Hiền", "Đào Ngọc Anh / Lê Thành Tỉnh"],
        ["GĐ 1: Khảo sát & CSDL", "B2 | 17/05", "Thiết kế các bảng Users, BienBans, ChiTietBienBans, ThanhPhanKiemTras, DinhLuongSuatAns và ChuKys.", "Xây dựng backend ASP.NET Core.", "Bùi Thanh Hiền", "Đào Ngọc Anh / Lê Thành Tỉnh"],
        ["GĐ 2: Backend API", "B1 | 19/05", "Hoàn thành đăng nhập, đăng ký, mã hóa mật khẩu BCrypt và sinh JWT token theo vai trò.", "Làm API biên bản và phân quyền.", "Bùi Thanh Hiền", "Đào Ngọc Anh / Lê Thành Tỉnh"],
        ["GĐ 2: Backend API", "B2 | 20/05", "Hoàn thành API tạo, lấy danh sách, xem chi tiết, cập nhật, gửi, duyệt và xóa biên bản.", "Kết nối frontend với API.", "Bùi Thanh Hiền", "Đào Ngọc Anh / Lê Thành Tỉnh"],
        ["GĐ 3: Frontend nhân viên", "B1 | 22/05", "Hoàn thành màn hình đăng nhập, điều hướng theo vai trò và dashboard nhân viên.", "Làm các form nhập biên bản.", "Bùi Thanh Hiền", "Đào Ngọc Anh / Lê Thành Tỉnh"],
        ["GĐ 3: Frontend nhân viên", "B2 | 23/05", "Hoàn thành 4 form: Cơ sở hạ tầng, Hồ sơ sổ sách, Vệ sinh ATTP và Suất ăn người bệnh.", "Bổ sung chữ ký điện tử và thông báo.", "Bùi Thanh Hiền", "Đào Ngọc Anh / Lê Thành Tỉnh"],
        ["GĐ 4: Admin & thống kê", "B1 | 24/05", "Hoàn thành dashboard Admin, danh sách biên bản chờ duyệt, xem chi tiết và chức năng duyệt.", "Hoàn thiện biểu đồ/thống kê.", "Bùi Thanh Hiền", "Đào Ngọc Anh / Lê Thành Tỉnh"],
        ["GĐ 4: Admin & thống kê", "B2 | 25/05", "Hoàn thành thống kê tổng quan, thống kê theo loại biên bản, theo ngày/tuần/tháng và tỷ lệ đạt.", "Kiểm thử toàn bộ luồng nghiệp vụ.", "Bùi Thanh Hiền", "Đào Ngọc Anh / Lê Thành Tỉnh"],
        ["GĐ 5: Hoàn thiện", "B1 | 26/05", "Kiểm thử đăng nhập, tạo biên bản, gửi biên bản, duyệt biên bản, thông báo và responsive giao diện.", "Chuẩn bị báo cáo và ảnh demo.", "Bùi Thanh Hiền", "Đào Ngọc Anh / Lê Thành Tỉnh"],
        ["GĐ 5: Hoàn thiện", "B2 | 26/05", "Tổng hợp thông tin dự án, mô tả chức năng, công nghệ, database, API và quy trình hoạt động.", "Đóng gói báo cáo, bổ sung ảnh demo giao diện.", "Bùi Thanh Hiền", "Đào Ngọc Anh / Lê Thành Tỉnh"],
    ])

    # Add demo placeholders below section 4.
    demo = next(
        paragraph for paragraph in doc.paragraphs
        if paragraph.text.strip().lower().startswith("4. demo")
    )
    current = demo
    placeholders = [
        "4.1 Màn hình đăng nhập: [Chèn ảnh demo màn hình LoginPage tại đây]",
        "4.2 Dashboard nhân viên: [Chèn ảnh demo EmployeeDashboard/EmployeeHome tại đây]",
        "4.3 Form lập biên bản: [Chèn ảnh demo 4 loại form biên bản tại đây]",
        "4.4 Chữ ký điện tử: [Chèn ảnh demo SignaturePadPage tại đây]",
        "4.5 Dashboard Admin và duyệt biên bản: [Chèn ảnh demo AdminDashboard tại đây]",
        "4.6 Thống kê biên bản: [Chèn ảnh demo biểu đồ/thống kê tại đây]",
    ]
    for text in placeholders:
        current = insert_paragraph_after(current, text)

    # Light formatting pass for all paragraphs.
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(("1.", "2.", "3.", "4.")):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.bold = True
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            if run.font.size is None:
                run.font.size = Pt(12)

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
